"""Generate QUANTUMEXE POS QC Test Plan Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tePr = cell._tc.get_or_add_tcPr()
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, fill="1F4E79"):
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)


def case_block(doc, case_id, title, role, preconditions, steps, expected):
    p = doc.add_paragraph()
    run = p.add_run(f"{case_id} — {title}")
    run.bold = True
    run.font.size = Pt(11)
    meta = doc.add_paragraph()
    meta.add_run(f"Role: {role}").font.size = Pt(10)
    if preconditions:
        doc.add_paragraph().add_run(f"Preconditions: {preconditions}").font.size = Pt(10)
    doc.add_paragraph().add_run("Steps:").bold = True
    for i, s in enumerate(steps, 1):
        bullet(doc, f"{i}. {s}")
    doc.add_paragraph().add_run("Expected result:").bold = True
    bullet(doc, expected)
    # Pass/Fail line
    pf = doc.add_paragraph()
    pf.add_run("Result:  [ ] Pass    [ ] Fail    [ ] Blocked     Tester: _____________     Date: _____________").font.size = Pt(9)
    doc.add_paragraph()


def build():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("QUANTUMEXE POS")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Quality Control (QC) Test Plan")
    r.bold = True
    r.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Full functional verification — all user roles & all modules\n"
        "Document version: 1.0  |  Date: 28 July 2026\n"
        "Application: QUANTUMEXE POS (quantumexe-pos)\n"
        "Prepared for: QC / QA Team"
    ).font.size = Pt(11)

    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project", "QUANTUMEXE POS"],
            ["Environments", "Cloud demo + Local / Desktop"],
            ["Cloud URL", "https://quantumexe-pos.vercel.app"],
            ["Local web", "http://localhost:5173 (API :4000)"],
            ["Desktop", "QUANTUMEXE-POS-Setup-*.exe (Windows)"],
            ["Build / version under test", "_______________ (fill before QC)"],
            ["QC lead", "_______________"],
            ["Start date", "_______________"],
            ["End date", "_______________"],
            ["Sign-off", "_______________"],
        ],
        [2.2, 5.0],
    )

    # 1. Purpose
    heading(doc, "1. Purpose", 1)
    para(
        doc,
        "This document defines end-to-end QC steps so testers can verify every major function "
        "for every user role before release. Mark Pass / Fail / Blocked on each case. "
        "Any Fail must include steps to reproduce, screenshot, and severity.",
    )

    heading(doc, "2. Severity guide", 1)
    add_table(
        doc,
        ["Severity", "Meaning", "Example"],
        [
            ["S1 – Critical", "Blocks core business; no workaround", "Cannot login, cannot checkout, data loss"],
            ["S2 – Major", "Core feature broken; workaround exists", "Receipt print fails, return rejects valid qty"],
            ["S3 – Minor", "UI/wording/partial issue", "Wrong label text, layout glitch"],
            ["S4 – Cosmetic", "Visual only", "Spacing, color contrast"],
        ],
    )

    heading(doc, "3. Test accounts (demo seed)", 1)
    para(doc, "Use these accounts after seed / demo shop is available. Change passwords if your environment differs.", bold=False)
    add_table(
        doc,
        ["Role", "Login (Contact / Username)", "Password", "Access summary"],
        [
            ["Master Admin", "master", "Master@123", "Shop registry only (/master)"],
            ["Shop Admin (Super Admin)", "0771234567", "123456", "Full shop modules (by shop type)"],
            ["Cashier", "0771111111", "123456", "Dashboard, Sales, Quotation*, Customers, POS"],
            ["Storekeeper", "0772222222", "123456", "Dashboard, Stock, GRN, Products, Supplier"],
        ],
    )
    bullet(doc, "Login field is phone/contact (not email), except Master Admin username.")
    bullet(doc, "*Quotation appears only if the shop type includes the quotation module.")

    heading(doc, "4. Environments & browsers", 1)
    bullet(doc, "Cloud: https://quantumexe-pos.vercel.app — hard refresh (Ctrl+F5) after deploy; re-login if session stale.")
    bullet(doc, "Local: Start API + web (or Start-POS-Dev.bat).")
    bullet(doc, "Desktop installer: recommended for offline / SQLite shops.")
    bullet(doc, "Browsers: Chrome and Edge (required for CD-7220 Web Serial). Also spot-check Firefox for core POS.")
    bullet(doc, "Hardware (optional): barcode scanner, thermal printer, second monitor, CD-7220 customer pole display.")

    heading(doc, "5. Shop types to cover", 1)
    para(doc, "Master Admin assigns shop type on approve. QC should smoke-test at least Clothing + one other type.")
    add_table(
        doc,
        ["Shop type", "Quotation", "GRN/Supplier", "Notes"],
        [
            ["Clothing / Fashion", "Yes", "Yes", "Brand shown; size/color important"],
            ["Restaurant / Cafe", "No", "Yes", "Quick sale; brand menu hidden"],
            ["Grocery / Mini mart", "Yes", "Yes", "Expire stock emphasis"],
            ["Pharmacy", "No", "Yes", "Expire stock emphasis"],
            ["Electronics", "Yes", "Yes", "Standard retail"],
            ["General retail", "Yes", "Yes", "Standard retail"],
        ],
    )

    heading(doc, "6. Role access matrix (smoke)", 1)
    add_table(
        doc,
        ["Module", "Master", "Admin", "Cashier", "Storekeeper"],
        [
            ["Master console", "Yes", "No", "No", "No"],
            ["Dashboard", "No", "Yes", "Yes", "Yes"],
            ["POS checkout", "No", "Yes", "Yes", "Yes*"],
            ["Sales / Returns", "No", "Yes", "Yes", "No (nav)"],
            ["Quotation", "No", "Yes*", "Yes*", "No"],
            ["Products / Catalog", "No", "Yes", "No", "Yes"],
            ["Stock / GRN", "No", "Yes", "No", "Yes"],
            ["Supplier", "No", "Yes*", "No", "Yes*"],
            ["Customers", "No", "Yes", "Yes", "No"],
            ["Manage Users", "No", "Yes", "No", "No"],
            ["Accounts", "No", "Yes", "No", "No"],
            ["Reports", "No", "Yes", "No", "No"],
            ["Settings / Backup", "No", "Yes", "No", "No"],
            ["Customer display page", "Public route", "Public", "Public", "Public"],
        ],
    )
    para(doc, "* = if module enabled for shop type. POS button is in top bar for authenticated shop users.")

    # ========== TEST CASES ==========
    heading(doc, "7. Detailed test cases", 1)

    heading(doc, "7.1 Authentication & access control", 2)

    case_block(
        doc, "AUTH-01", "Valid shop Admin login",
        "Admin", "Demo shop active",
        ["Open Sign In", "Enter 0771234567 / 123456", "Submit"],
        "Dashboard loads; sidebar shows Admin modules; no redirect to pending-access.",
    )
    case_block(
        doc, "AUTH-02", "Valid Cashier login",
        "Cashier", "User active",
        ["Login as 0771111111 / 123456", "Inspect sidebar"],
        "Sees Dashboard, Sales, Customers, POS; does NOT see Settings, Stock, Products, Users.",
    )
    case_block(
        doc, "AUTH-03", "Valid Storekeeper login",
        "Storekeeper", "User active",
        ["Login as 0772222222 / 123456", "Inspect sidebar"],
        "Sees Stock, GRN, Products, Supplier; does NOT see Sales / Customers / Settings.",
    )
    case_block(
        doc, "AUTH-04", "Master Admin login",
        "MasterAdmin", None,
        ["Login as master / Master@123"],
        "Redirected to /master shop registry. Cannot use shop POS layout.",
    )
    case_block(
        doc, "AUTH-05", "Invalid credentials",
        "Any", None,
        ["Enter wrong password", "Submit"],
        "Login fails with clear error (401). No session created.",
    )
    case_block(
        doc, "AUTH-06", "Pending / revoked shop gate",
        "Admin of pending shop", "Shop status pending or revoked",
        ["Login with that shop’s admin"],
        "User lands on pending-access; POS APIs blocked until Master approves.",
    )
    case_block(
        doc, "AUTH-07", "Logout / session",
        "Admin", "Logged in",
        ["Use logout", "Try to open /dashboard directly"],
        "Redirected to Sign In. Protected routes require re-login.",
    )

    heading(doc, "7.2 Master Admin", 2)
    case_block(
        doc, "MST-01", "List shops & open detail",
        "MasterAdmin", "Logged in as master",
        ["Open Master console", "Select a shop", "Review Firebase / Payment / Security tabs"],
        "Shop details load; tabs switch without error.",
    )
    case_block(
        doc, "MST-02", "Approve shop with shop type",
        "MasterAdmin", "Pending shop exists",
        ["Open shop", "Choose shop type (e.g. Clothing)", "Approve / mark paid as per UI"],
        "Shop becomes active; type-based modules/features applied; Admin can enter POS.",
    )
    case_block(
        doc, "MST-03", "Approve without shop type (negative)",
        "MasterAdmin", "Pending shop",
        ["Attempt approve without selecting type"],
        "Blocked with validation error (400).",
    )
    case_block(
        doc, "MST-04", "Change shop type",
        "MasterAdmin", "Active shop",
        ["Change type from Clothing to Restaurant", "Login as shop Admin", "Check Quotation / Brand menus"],
        "Nav matches new type (Restaurant: no Quotation; brand may be hidden).",
    )
    case_block(
        doc, "MST-05", "Revoke access",
        "MasterAdmin", "Active shop",
        ["Revoke shop", "Login as shop Admin"],
        "Pending-access / blocked until restored.",
    )
    case_block(
        doc, "MST-06", "Reset Super Admin password",
        "MasterAdmin", "Active shop",
        ["Reset shop admin password", "Login with new password"],
        "New password works; old password fails.",
    )
    case_block(
        doc, "MST-07", "Firebase connect (if used)",
        "MasterAdmin", "Shop Firebase credentials ready",
        ["Paste Project ID + SA email + private key (or JSON)", "Save & provision"],
        "Connection succeeds; shop data routes to dedicated project (or clear error if invalid).",
    )

    heading(doc, "7.3 Dashboard", 2)
    case_block(
        doc, "DASH-01", "Dashboard loads for each role",
        "Admin / Cashier / Storekeeper", "Logged in",
        ["Open Dashboard", "Check cards / charts / recent data"],
        "Page loads without error for each role; no blank crash.",
    )

    heading(doc, "7.4 Point of Sale (POS)", 2)
    case_block(
        doc, "POS-01", "Add products by click",
        "Cashier", "Stocked products exist",
        ["Open POS", "Click a product card", "Increase qty"],
        "Item appears in cart; line total updates; qty cannot exceed stock.",
    )
    case_block(
        doc, "POS-02", "Barcode scan / entry",
        "Cashier", "Known barcode in catalog",
        ["Enter barcode in scan field", "Submit"],
        "Matching product added. Unknown barcode shows “not found”.",
    )
    case_block(
        doc, "POS-03", "Checkout cash sale",
        "Cashier", "Cart has items",
        ["Select payment Cash", "Enter paid amount ≥ total", "Checkout"],
        "Invoice created; success message with invoice no; stock reduced.",
    )
    case_block(
        doc, "POS-04", "Auto print receipt",
        "Cashier", "Auto print checked; pop-ups allowed",
        ["Complete checkout with Auto print ON"],
        "Print dialog / receipt window opens with shop header, items (size/color), totals, barcode if enabled.",
    )
    case_block(
        doc, "POS-05", "Auto print OFF",
        "Cashier", None,
        ["Uncheck Auto print", "Checkout"],
        "Sale succeeds; no automatic print dialog.",
    )
    case_block(
        doc, "POS-06", "Insufficient stock (negative)",
        "Cashier", "Product with low stock",
        ["Try to checkout qty greater than available (API/force if UI caps)"],
        "Sale rejected with insufficient stock message; stock unchanged.",
    )
    case_block(
        doc, "POS-07", "Empty cart checkout (negative)",
        "Cashier", None,
        ["Clear cart", "Click Checkout"],
        "Blocked with “Cart is empty” (or equivalent).",
    )
    case_block(
        doc, "POS-08", "Display name shows size & color",
        "Cashier", "Variant with size + color",
        ["Add variant to cart", "Inspect name on POS card/cart"],
        "Name includes Size and Color (not only product name).",
    )

    heading(doc, "7.5 Sales, invoices & returns", 2)
    case_block(
        doc, "SAL-01", "Manage Invoice list",
        "Admin / Cashier", "Prior sales exist",
        ["Open Sales → Manage Invoice", "Search / filter", "Open invoice detail"],
        "Invoices list correctly; detail shows items and totals.",
    )
    case_block(
        doc, "SAL-02", "Reprint receipt from invoice",
        "Admin / Cashier", "Invoice exists",
        ["From Manage Invoice, print receipt"],
        "Receipt matches sale; settings (header/footer/logo/barcode) respected.",
    )
    case_block(
        doc, "SAL-03", "User Sales view",
        "Admin / Cashier", None,
        ["Open Manage User Sales"],
        "Per-user sales summary loads without error.",
    )
    case_block(
        doc, "SAL-04", "Partial return",
        "Admin / Cashier", "Invoice with multiple lines / qty > 1",
        [
            "Open Return History",
            "Search invoice number",
            "Set return qty on ONE line only (less than remaining)",
            "Process Return",
        ],
        "Only selected qty returned; refund total correct; stock increased for that variant; remaining qty updates.",
    )
    case_block(
        doc, "SAL-05", "Second partial return on same invoice",
        "Admin / Cashier", "Invoice already partially returned",
        ["Search same invoice", "Return more qty within remaining"],
        "Allowed up to remaining. Over-return shows error: cannot return more than left.",
    )
    case_block(
        doc, "SAL-06", "Select all remaining",
        "Admin / Cashier", "Invoice with remaining qty",
        ["Click Select all remaining", "Process Return"],
        "All remaining lines returned; invoice has no remaining qty afterward.",
    )
    case_block(
        doc, "SAL-07", "Exchange flow",
        "Admin / Cashier", "Invoice with remaining qty",
        ["Select return qty", "Click Exchange → POS"],
        "Return succeeds; navigates to POS for replacement sale.",
    )
    case_block(
        doc, "SAL-08", "Return unknown invoice (negative)",
        "Admin / Cashier", None,
        ["Search invalid invoice number"],
        "Not found / clear error; no crash.",
    )

    heading(doc, "7.6 Products & catalog", 2)
    case_block(
        doc, "PRD-01", "Create product with size & color variants",
        "Admin / Storekeeper", None,
        [
            "Products → Create Product",
            "Fill Basic → Inventory → Variations",
            "Add Size and Color on variant(s)",
            "Save",
        ],
        "Product appears in Product List; variants store size & color.",
    )
    case_block(
        doc, "PRD-02", "Manage Category / Brand / Unit / Product Type",
        "Admin / Storekeeper", None,
        ["Create and edit each lookup", "Use quick-add from Create Product if available"],
        "CRUD works; new values selectable on product form.",
    )
    case_block(
        doc, "PRD-03", "Deactivate product",
        "Admin / Storekeeper", "Existing product",
        ["Deactivate from Product List", "Check Deactivated list", "Confirm not sold as normal stock"],
        "Deactivated products listed separately; not offered as normal sellable stock.",
    )
    case_block(
        doc, "PRD-04", "Print product labels (single)",
        "Admin / Storekeeper", "Product with barcode",
        ["Product List → tag/print icon on one product"],
        "Label sheet opens with shop name, product name, size/color, price, barcode.",
    )
    case_block(
        doc, "PRD-05", "Print labels (page)",
        "Admin / Storekeeper", "Product list has rows",
        ["Click Print labels (page)"],
        "Labels generated for variants on current page; empty selection handled gracefully.",
    )

    heading(doc, "7.7 Stock & GRN", 2)
    case_block(
        doc, "STK-01", "Stock list & filters",
        "Admin / Storekeeper", "Stock exists",
        ["Open Stock List", "Search / filter", "Open Low / Out / Damaged / Expire views"],
        "Each view loads; quantities match expectations after sales/returns.",
    )
    case_block(
        doc, "STK-02", "Create GRN",
        "Admin / Storekeeper", "Supplier + products exist; GRN module on",
        ["GRN → Create GRN", "Add lines", "Save", "Open GRN List"],
        "GRN saved; stock increased accordingly.",
    )
    case_block(
        doc, "STK-03", "Damaged stock adjust (negative qty)",
        "Admin / Storekeeper", "Stock available",
        ["Mark damaged more than available qty"],
        "Rejected with insufficient stock error.",
    )

    heading(doc, "7.8 Quotation", 2)
    case_block(
        doc, "QUO-01", "Create & list quotation",
        "Admin / Cashier", "Shop type includes quotation",
        ["Create quotation with customer + items", "Save", "Open Quotation List", "Edit if available"],
        "Quotation saved with correct total; list/edit work.",
    )
    case_block(
        doc, "QUO-02", "Quotation hidden when module off",
        "Admin", "Restaurant or Pharmacy shop",
        ["Inspect sidebar"],
        "Quotation menu not shown.",
    )

    heading(doc, "7.9 Customers, suppliers, users", 2)
    case_block(
        doc, "CUS-01", "Manage customers",
        "Admin / Cashier", None,
        ["Create customer", "Edit", "Search", "Select customer on POS"],
        "Customer saved; appears on POS customer dropdown.",
    )
    case_block(
        doc, "SUP-01", "Manage suppliers & company",
        "Admin / Storekeeper", "Supplier module on",
        ["Create supplier", "Manage company", "Link to GRN / payments if available"],
        "Supplier CRUD works; usable in GRN.",
    )
    case_block(
        doc, "USR-01", "Manage users — create",
        "Admin", None,
        ["Manage Users → add Cashier or Storekeeper", "Login as new user"],
        "New user can login with correct role nav.",
    )
    case_block(
        doc, "USR-02", "Delete user rules",
        "Admin", "Multiple admins / self logged in",
        ["Try delete own account", "Try delete last active admin"],
        "Self-delete and last-admin delete blocked with clear message.",
    )
    case_block(
        doc, "USR-03", "Delete other user",
        "Admin", "Extra non-critical user exists",
        ["Delete that user", "Try login as deleted user"],
        "User removed; login fails.",
    )

    heading(doc, "7.10 Accounts, employees, reports", 2)
    case_block(
        doc, "ACC-01", "Accounts overview",
        "Admin", None,
        ["Open Accounts"],
        "Page loads (sessions / cash overview) without crash.",
    )
    case_block(
        doc, "EMP-01", "Employee module (direct URL)",
        "Admin", None,
        ["Open /employee/manage-employee", "Create employee", "Attendance mark", "Salary if available"],
        "CRUD / attendance / salary pages function; note if not linked in sidebar.",
    )
    case_block(
        doc, "RPT-01", "Sales financial report",
        "Admin", "Sales data exists",
        ["Open Sales Financial Report", "Apply date filters"],
        "Totals reconcile with sample invoices for the period.",
    )
    case_block(
        doc, "RPT-02", "Inventory / product report",
        "Admin", "Stock exists",
        ["Open Inventory Product Report"],
        "Stock quantities/values look correct; export/print if present.",
    )
    case_block(
        doc, "RPT-03", "Tax / employee / quotation reports",
        "Admin", None,
        ["Open each report from Reports area"],
        "Each report loads; empty state is friendly when no data.",
    )

    heading(doc, "7.11 Settings, print & customer display", 2)
    case_block(
        doc, "SET-01", "POS settings save",
        "Admin", None,
        ["Settings → POS", "Change shop display name, currency, toggles", "Save", "Reload"],
        "Values persist.",
    )
    case_block(
        doc, "SET-02", "Print settings drive receipt",
        "Admin", None,
        [
            "Settings → Print",
            "Set header/footer, thermal/A4, logo/barcode toggles",
            "Save",
            "Print a test receipt from POS or invoice",
        ],
        "Receipt reflects settings.",
    )
    case_block(
        doc, "SET-03", "Open customer display window",
        "Admin", "Pop-ups allowed; Customer Display enabled",
        [
            "Settings → Display → Open customer display window",
            "Drag window to second monitor",
            "On POS, add items to cart",
        ],
        "Customer screen shows live items, prices, subtotal/discount/total; thank-you after sale.",
    )
    case_block(
        doc, "SET-04", "Customer display disabled",
        "Admin", None,
        ["Turn off Customer Display sync", "Change POS cart"],
        "Customer window does not update (or stays idle per design).",
    )
    case_block(
        doc, "SET-05", "CD-7220 pole display connect",
        "Admin / Cashier", "Chrome/Edge + CD-7220 on USB/COM",
        [
            "Settings or POS → Connect CD-7220",
            "Select serial port",
            "Add cart items",
            "Checkout",
        ],
        "Pole shows item/total lines; thank-you after sale; disconnect works. Non-Chrome shows unsupported message.",
    )
    case_block(
        doc, "SET-06", "License & About",
        "Admin", None,
        ["Settings → License", "Settings → About"],
        "License status visible; About shows current app version (not stale DB version).",
    )
    case_block(
        doc, "SET-07", "Connection / sync center",
        "Admin", None,
        ["Settings → Connection", "Review online/offline indicators"],
        "Status accurate; no unexplained errors.",
    )

    heading(doc, "7.12 Backup", 2)
    case_block(
        doc, "BKP-01", "Create backup",
        "Admin", None,
        ["Open Backup", "Create backup", "Confirm file/list entry appears"],
        "Backup created successfully.",
    )
    case_block(
        doc, "BKP-02", "Delete backup (if supported)",
        "Admin", "Backup exists",
        ["Delete a backup entry"],
        "Removed from list without corrupting live DB.",
    )

    heading(doc, "8. Cross-cutting regression checklist", 1)
    add_table(
        doc,
        ["ID", "Check", "Pass?"],
        [
            ["X-01", "Hard refresh after deploy; re-login works", "[ ]"],
            ["X-02", "Mobile / narrow width: POS usable enough for smoke", "[ ]"],
            ["X-03", "Currency / LKR formatting consistent", "[ ]"],
            ["X-04", "No console-breaking JS errors on main pages (spot-check DevTools)", "[ ]"],
            ["X-05", "API error messages shown in UI (not silent fail)", "[ ]"],
            ["X-06", "After return, POS stock list refreshes correctly", "[ ]"],
            ["X-07", "Clothing: size + color on POS, receipt, and labels", "[ ]"],
            ["X-08", "Restaurant: Quotation hidden; quick-sale behavior OK", "[ ]"],
            ["X-09", "Pharmacy/Grocery: expire stock screens meaningful", "[ ]"],
            ["X-10", "Master cannot access shop routes; shop user cannot access /master", "[ ]"],
        ],
    )

    heading(doc, "9. Recommended QC execution order", 1)
    bullet(doc, "Day 1: Auth + Master Admin approve clothing shop + role access matrix.")
    bullet(doc, "Day 1–2: Catalog (size/color) → Stock/GRN → POS sales → Receipt/labels.")
    bullet(doc, "Day 2: Partial returns / exchange → Customer display + pole (if hardware).")
    bullet(doc, "Day 2–3: Quotation, Customers, Users, Reports, Settings, Backup.")
    bullet(doc, "Day 3: Second shop type smoke (Restaurant or Pharmacy) + regression X-01…X-10.")
    bullet(doc, "Final: Bug triage meeting; only S1/S2 open bugs block release.")

    heading(doc, "10. Defect log template", 1)
    add_table(
        doc,
        ["Bug ID", "Case ID", "Severity", "Summary", "Status"],
        [
            ["BUG-001", "", "S1/S2/S3/S4", "", "Open/Fixed"],
            ["BUG-002", "", "", "", ""],
            ["BUG-003", "", "", "", ""],
            ["BUG-004", "", "", "", ""],
            ["BUG-005", "", "", "", ""],
        ],
    )
    para(
        doc,
        "For each bug attach: environment, role, exact steps, expected vs actual, screenshot, and invoice/product IDs if relevant.",
    )

    heading(doc, "11. Sign-off", 1)
    add_table(
        doc,
        ["Role", "Name", "Signature", "Date"],
        [
            ["QC Tester", "", "", ""],
            ["QC Lead", "", "", ""],
            ["Product / Dev", "", "", ""],
            ["Release approval", "", "", ""],
        ],
    )

    para(
        doc,
        "End of document — QUANTUMEXE POS QC Test Plan v1.0. "
        "Send this file to QC as-is; fill Build/version and dates before starting.",
        bold=False,
    )

    out = Path(r"C:\Users\p\Projects\reox-pos-clone\docs\QUANTUMEXE-POS-QC-Test-Plan.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    # Also copy to Desktop for easy find
    desk = Path.home() / "Desktop" / "QUANTUMEXE-POS-QC-Test-Plan.docx"
    try:
        doc.save(desk)
    except Exception:
        pass
    print(str(out))
    if desk.exists():
        print(str(desk))


if __name__ == "__main__":
    build()
