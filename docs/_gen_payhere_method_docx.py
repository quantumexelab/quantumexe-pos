"""Generate QUANTUMEXE PayHere SaaS Subscription Method handoff Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, fill="0B3D2E"):
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(col_widths):
                    row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.size = Pt(10)


def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("QUANTUMEXE — PayHere SaaS Subscription Method")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(11, 61, 46)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(
        "Handoff guide for AI agents / developers — reuse this pattern when adding subscriptions to other QUANTUMEXE systems"
    )
    s.font.size = Pt(10)
    s.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "Reference implementation: QUANTUMEXE POS  ·  Domain: quantumexe.com  ·  Updated: August 2026"
    )
    m.font.size = Pt(9)

    doc.add_heading("1. Purpose of this document", level=1)
    para(
        doc,
        "Give this document to another agent or developer when adding a PayHere subscription "
        "to a new QUANTUMEXE product. Follow the same architecture that already works for POS — "
        "do not invent a new checkout flow that posts directly from *.vercel.app.",
    )
    para(doc, "Primary goals:", bold=True)
    bullet(doc, "Shop / tenant / customer pays QUANTUMEXE via PayHere (LKR).")
    bullet(doc, "PayHere settles to QUANTUMEXE merchant bank — app never stores card numbers.")
    bullet(doc, "Webhook activates or extends access automatically.")
    bullet(doc, "Avoid PayHere “Unauthorized payment request” on production.")

    doc.add_heading("2. Critical PayHere constraint (read first)", level=1)
    para(
        doc,
        "PayHere Sandbox Integrations does NOT allow subdomains. Registering domains like "
        "pos.quantumexe.com or anything.vercel.app fails with: “Sub Domains not allowed or Invalid Domain name”.",
        bold=True,
    )
    bullet(doc, "Allowed examples: localhost (dev only), quantumexe.com (apex).")
    bullet(doc, "Not allowed: pos.quantumexe.com, app.example.com, project.vercel.app.")
    para(
        doc,
        "Because of this, the browser must open / POST the PayHere checkout form from the "
        "registered apex domain (Referer = quantumexe.com), even if the product UI lives on a subdomain.",
    )

    doc.add_heading("3. Working architecture (POS reference)", level=1)
    para(doc, "Hosts", bold=True)
    add_table(
        doc,
        ["Role", "URL", "Hosting"],
        [
            ["Company / marketing site + PayHere bridge origin", "https://quantumexe.com", "Vercel project: our_s_company-main (repo: quantumexelab/quantum_web)"],
            ["POS application UI + API", "https://pos.quantumexe.com", "Vercel project: quantumexe-pos (repo: quantumexelab/quantumexe-pos)"],
            ["Legacy / alternate UI URL", "https://quantumexe-pos.vercel.app", "Same POS project — do NOT use as PayHere Integrations domain"],
            ["Local development", "http://localhost:5173 (+ API :4000)", "npm run dev — PayHere domain can be localhost"],
        ],
        col_widths=[1.6, 2.2, 3.0],
    )

    para(doc, "End-to-end payment flow", bold=True)
    numbered(doc, "User is logged into the product (e.g. pos.quantumexe.com) as Admin / owner.")
    numbered(doc, "UI calls POST /api/billing/checkout (authenticated) with plan interval (monthly|annual) and optional recurring flag.")
    numbered(doc, "API builds PayHere form fields + MD5 hash using Merchant ID + Merchant Secret (server-side only).")
    numbered(doc, "API signs a short-lived JWT bridge token (audience: payhere-bridge, ~15 minutes) containing { action, fields }.")
    numbered(doc, "API returns bridgeUrl = https://quantumexe.com/api/billing/bridge?t=<JWT> (plus raw fields for local fallback).")
    numbered(doc, "Browser navigates to bridgeUrl on apex (window.location.assign).")
    numbered(doc, "Company site Vercel rewrite proxies /api/billing/* → https://pos.quantumexe.com/api/billing/* (URL bar stays quantumexe.com).")
    numbered(doc, "Bridge HTML page auto-POSTs the form to sandbox.payhere.lk (or live). Referer = quantumexe.com → PayHere accepts.")
    numbered(doc, "User pays on PayHere. notify_url hits POS webhook. return_url / cancel_url go via apex paths that rewrite back to POS.")
    numbered(doc, "Webhook verifies md5sig, then applySubscriptionPayment() extends nextDueAt / sets status active.")

    doc.add_heading("4. Company-site proxy (required for apex)", level=1)
    para(
        doc,
        "Repo quantumexelab/quantum_web contains vercel.json rewrites. Do not remove these when editing the company site:",
    )
    code_block(
        doc,
        '{\n'
        '  "rewrites": [\n'
        '    {\n'
        '      "source": "/api/billing/:path*",\n'
        '      "destination": "https://pos.quantumexe.com/api/billing/:path*"\n'
        '    },\n'
        '    {\n'
        '      "source": "/pos-return",\n'
        '      "destination": "https://pos.quantumexe.com/setting?tab=license&billing=return"\n'
        '    },\n'
        '    {\n'
        '      "source": "/pos-cancel",\n'
        '      "destination": "https://pos.quantumexe.com/setting?tab=license&billing=cancel"\n'
        '    }\n'
        '  ]\n'
        '}',
    )
    para(
        doc,
        "For a NEW product on another subdomain (e.g. salon.quantumexe.com), either: "
        "(A) add more rewrite destinations with product-specific paths, or "
        "(B) keep one shared billing API on POS/control plane and pass productId in custom fields.",
    )

    doc.add_heading("5. Environment variables", level=1)
    para(doc, "Set on the product API host (POS: Vercel → quantumexe-pos → Environment Variables → Production):", bold=True)
    add_table(
        doc,
        ["Variable", "Example / notes"],
        [
            ["PAYHERE_MODE", "sandbox | live"],
            ["PAYHERE_MERCHANT_ID", "From PayHere Integrations (e.g. 1237256)"],
            ["PAYHERE_MERCHANT_SECRET", "Secret for apex domain quantumexe.com ONLY (not localhost secret). Keep == padding."],
            ["PAYHERE_CHECKOUT_BASE", "https://quantumexe.com"],
            ["PAYHERE_RETURN_BASE", "https://quantumexe.com"],
            ["PAYHERE_BRIDGE_PATH", "/api/billing/bridge"],
            ["PAYHERE_RETURN_PATH", "/pos-return"],
            ["PAYHERE_CANCEL_PATH", "/pos-cancel"],
            ["PUBLIC_API_BASE", "https://pos.quantumexe.com  (webhook notify host)"],
            ["PUBLIC_WEB_BASE", "https://pos.quantumexe.com"],
            ["PAYHERE_AMOUNT_MONTHLY", "Optional, default 2000"],
            ["PAYHERE_AMOUNT_ANNUAL", "Optional, default 20000"],
            ["PAYHERE_DISABLE_RECURRING", "Set 1 to force one-time only"],
            ["JWT_SECRET", "Must match across verify of bridge JWT (same API that signs & verifies)"],
        ],
        col_widths=[2.4, 4.4],
    )
    para(
        doc,
        "After changing env vars: Redeploy the product. If company-site vercel.json changed: redeploy our_s_company-main too.",
    )

    doc.add_heading("6. Code map (POS reference files)", level=1)
    add_table(
        doc,
        ["File", "Responsibility"],
        [
            ["apps/api/src/billing/payhere.ts", "Hash, sanitize secret, plans, checkout fields, bridge JWT, config diagnostics"],
            ["apps/api/src/routes-billing.ts", "GET /billing/plans, POST /billing/checkout, GET /billing/bridge, GET /billing/bridge-json, POST /billing/webhook"],
            ["apps/api/src/master/shopRegistry.ts", "applySubscriptionPayment / markSubscriptionFailed — extend tenant access"],
            ["apps/web/src/pages/Settings.tsx", "License tab — plan select, startCheckout → bridgeUrl"],
            ["apps/web/src/pages/PendingAccess.tsx", "Pay gate for overdue/pending shops"],
            ["vercel.json (POS)", "Default PUBLIC_* and PAYHERE_* for deploy"],
            ["quantum_web/vercel.json", "Apex rewrites to POS billing + return paths"],
        ],
        col_widths=[2.6, 4.2],
    )

    doc.add_heading("7. API contracts to reimplement", level=1)
    para(doc, "POST /api/billing/checkout (auth required)", bold=True)
    bullet(doc, "Body: { interval: \"monthly\"|\"annual\", recurring?: boolean }")
    bullet(doc, "Response data must include: action (PayHere URL), fields (form map), bridgeUrl (preferred), orderId, amount")
    bullet(doc, "Put tenant id in PayHere custom_1; plan id in custom_2")

    para(doc, "GET /api/billing/bridge?t=JWT (public)", bold=True)
    bullet(doc, "Verify JWT → render HTML form → auto-submit to PayHere")
    bullet(doc, "Without t: return “Missing checkout token” (this is healthy)")

    para(doc, "POST /api/billing/webhook (public, form-urlencoded)", bold=True)
    bullet(doc, "Verify merchant_id + md5sig")
    bullet(doc, "status_code 2 (and recurring success codes as implemented) → activate / extend access")
    bullet(doc, "Failure codes → mark overdue / failed")
    bullet(doc, "Always respond so PayHere stops retrying appropriately")

    doc.add_heading("8. Frontend checkout rule", level=1)
    para(doc, "Preferred:", bold=True)
    code_block(doc, 'if (data.bridgeUrl) {\n  window.location.assign(data.bridgeUrl);\n  return;\n}\n// local/dev fallback only: submit form from current origin')
    para(
        doc,
        "Never teach agents to only POST from *.vercel.app with the apex merchant secret — that causes Unauthorized.",
    )

    doc.add_heading("9. How to reuse for another QUANTUMEXE system", level=1)
    para(doc, "Option A — Shared billing control plane (recommended)", bold=True)
    numbered(doc, "Keep checkout + webhook on one central API (or extend POS billing with productCode).")
    numbered(doc, "custom_1 = tenantId, custom_2 = plan, add custom field or order_id prefix for productCode.")
    numbered(doc, "Webhook routes activation to the correct product database / license table.")
    numbered(doc, "Reuse quantumexe.com bridge rewrites (same apex).")

    para(doc, "Option B — Copy pattern into the new app", bold=True)
    numbered(doc, "Copy payhere.ts + billing routes + License UI patterns.")
    numbered(doc, "Point PUBLIC_API_BASE at the new app API.")
    numbered(doc, "Add company-site rewrites for that API (or path-prefix /api/<product>/billing).")
    numbered(doc, "Register the SAME apex quantumexe.com in PayHere (one secret) OR a different apex if you buy another domain.")
    numbered(doc, "Never register the new app’s vercel.app URL in PayHere.")

    para(doc, "Data model minimum fields on tenant/license record", bold=True)
    bullet(doc, "status: pending | active | overdue | revoked")
    bullet(doc, "billingPlan / billingInterval")
    bullet(doc, "nextDueAt, lastPaidAt")
    bullet(doc, "payherePaymentId, payhereSubscriptionId (optional)")
    bullet(doc, "lastBillingAmount")

    doc.add_heading("10. PayHere merchant setup checklist", level=1)
    numbered(doc, "Create sandbox merchant at https://sandbox.payhere.lk (live later at payhere.lk).")
    numbered(doc, "Integrations → Add Domain/App → Domain Name = quantumexe.com (no https://).")
    numbered(doc, "Copy Merchant ID + Merchant Secret for that row (not the localhost row).")
    numbered(doc, "Enable Recurring / Subscriptions if auto-renew is required.")
    numbered(doc, "Put secret in Vercel; Redeploy.")
    numbered(doc, "Live go-live: new live merchant credentials, PAYHERE_MODE=live, bank payout setup in PayHere.")

    doc.add_heading("11. Local development", level=1)
    bullet(doc, "PayHere Integrations domain: localhost")
    bullet(doc, "apps/api/.env: PAYHERE_* + PUBLIC_API_BASE=http://localhost:4000 + PUBLIC_WEB_BASE=http://localhost:5173")
    bullet(doc, "Leave PAYHERE_CHECKOUT_BASE empty so local form post works, OR point bridge at local API")
    bullet(doc, "npm run dev → http://localhost:5173")
    bullet(doc, "If DB schema drifts: prisma db push --accept-data-loss && db:seed (local only)")

    doc.add_heading("12. Known failures & fixes", level=1)
    add_table(
        doc,
        ["Symptom", "Cause", "Fix"],
        [
            ["Unauthorized payment request", "Form posted from unregistered host (vercel.app / subdomain) or wrong secret", "Use apex bridge + secret for quantumexe.com"],
            ["Sub Domains not allowed", "Tried to register pos.* or vercel.app in PayHere", "Register apex only"],
            ["PayHere not configured / MISSING", "Env missing on that server", "Set PAYHERE_MERCHANT_ID/SECRET; redeploy"],
            ["Missing checkout token", "Opened /api/billing/bridge without ?t=", "Normal — only open via Pay button"],
            ["Company homepage broken", "Removed apex rewrites incorrectly", "Restore quantum_web vercel.json rewrites; do not point whole apex away from company project"],
            ["Webhook not activating", "notify_url not public / wrong PUBLIC_API_BASE / bad md5sig", "Use https://pos…/api/billing/webhook; verify secret; check logs"],
            ["Login timeout locally", "SQLite schema missing columns", "db push + seed"],
        ],
        col_widths=[2.0, 2.4, 2.4],
    )

    doc.add_heading("13. Security rules for agents", level=1)
    bullet(doc, "Never commit Merchant Secret to git or paste into client-side JS.")
    bullet(doc, "Always compute PayHere hash on the server.")
    bullet(doc, "Always verify webhook md5sig before granting access.")
    bullet(doc, "Bridge JWT must expire quickly (≈15m) and use JWT_SECRET.")
    bullet(doc, "Rotate secrets if they were pasted into chat logs.")

    doc.add_heading("14. Prompt snippet — paste to another agent", level=1)
    para(doc, "Copy/paste the following instruction block:", bold=True)
    code_block(
        doc,
        "Add PayHere SaaS subscriptions using the QUANTUMEXE method documented in\n"
        "docs/QUANTUMEXE-PayHere-Subscription-Method.docx (same repo / this handoff).\n\n"
        "Hard rules:\n"
        "1) Do NOT register *.vercel.app or product subdomains in PayHere — apex quantumexe.com only.\n"
        "2) Checkout must redirect to https://quantumexe.com/api/billing/bridge?t=<jwt>\n"
        "   (company site proxies /api/billing/* to the product API).\n"
        "3) Hash + secret stay server-side; webhook verifies md5sig then extends tenant access.\n"
        "4) Reuse pattern from apps/api/src/billing/payhere.ts and routes-billing.ts in quantumexe-pos.\n"
        "5) Keep company homepage unchanged except vercel.json billing rewrites.\n"
        "6) Env: PAYHERE_CHECKOUT_BASE/RETURN_BASE=https://quantumexe.com;\n"
        "   PUBLIC_API_BASE=<product https api host>; PAYHERE_MERCHANT_SECRET for apex.\n",
    )

    doc.add_heading("15. Quick verification checklist", level=1)
    bullet(doc, "https://quantumexe.com/api/billing/bridge → “Missing checkout token” (proxy OK)")
    bullet(doc, "Product License → Pay → lands on sandbox.payhere.lk checkout (not Unauthorized)")
    bullet(doc, "Complete sandbox payment → webhook → tenant status active / nextDueAt extended")
    bullet(doc, "https://quantumexe.com homepage still shows company marketing site")

    out = Path(__file__).resolve().parent / "QUANTUMEXE-PayHere-Subscription-Method.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
